#!/usr/bin/env python3
"""Rendu Word (.docx) du gabarit à partir du R Markdown (source unique : Gabarit/gabarit_snowballing.Rmd).

Reproduit la mise en forme exigée sans R ni LaTeX : Times New Roman 12 pt, interligne double, texte justifié,
pages numérotées (à partir de la page suivant la page titre), identifiants en gras et en majuscules en haut à
droite de chaque page, nombre de mots sur la page titre (page titre et bibliographie exclues), aucun nom.
Produit Gabarit/gabarit_snowballing.docx ; conversion PDF facultative avec LibreOffice.
"""
import os, re, datetime
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RMD = os.path.join(ROOT, "Gabarit", "gabarit_snowballing.Rmd")
GRILLE = os.path.join(ROOT, "Data", "grille_snowballing.csv")
PNG = os.path.join(ROOT, "Graphiques", "lectures_snowballing_top30_court.png")
PNG_RESEAU = os.path.join(ROOT, "Graphiques", "reseau_citations_top30.png")
OUT = os.path.join(ROOT, "Gabarit", "gabarit_snowballing.docx")
IDUL = "537380130 / MADUP212"
W_CIT, W_YEAR, W_PERT, N_TOP = 0.35, 0.25, 0.40, 30

# ---------- 1. Lecture du Rmd : texte, titres, listes, emplacement des figures ----------
lines = open(RMD, encoding="utf-8").read().split("\n")
sep = [i for i, l in enumerate(lines) if l.strip() == "---"]
body = lines[sep[1] + 1:] if len(sep) >= 2 else lines

blocks = []          # (type, contenu)
in_chunk, chunk_label, in_title, in_comment = False, None, False, False
para = []
def flush():
    global para
    if para:
        blocks.append(("p", " ".join(para).strip()))
        para = []
for l in body:
    s = l.strip()
    if s.startswith("```"):
        if not in_chunk:
            in_chunk = True
            m = re.match(r"```\{r\s*([A-Za-z0-9_]*)", s)
            chunk_label = m.group(1) if m else ""
            flush()
            if chunk_label in ("graphique", "reseau"):
                blocks.append(("figure:" + chunk_label, re.search(r'fig\.cap="([^"]*)"', s).group(1)))
            elif chunk_label == "tableau":
                blocks.append(("tableau", ""))
            elif chunk_label == "bibliographie":
                blocks.append(("bibliographie", ""))
        else:
            in_chunk = False
        continue
    if in_chunk:
        continue
    if "\\begin{titlepage}" in s: in_title = True; continue
    if "\\end{titlepage}" in s: in_title = False; continue
    if in_title: continue
    if s.startswith("<!--"): in_comment = True
    if in_comment:
        if "-->" in s: in_comment = False
        continue
    if s.startswith("#"):
        flush()
        level = len(s) - len(s.lstrip("#"))
        blocks.append((f"h{level}", re.sub(r"\s*\{-\}$", "", s.lstrip("#").strip())))
    elif s.startswith("- "):
        flush(); blocks.append(("li", s[2:].strip()))
    elif s == "":
        flush()
    elif s.startswith("\\"):
        continue                                  # commande LaTeX (\newpage…) : ignorée dans le rendu Word
    else:
        para.append(s)
flush()

# ---------- 2. Grille, indice et top 30 (même formule que le Rmd) ----------
g = pd.read_csv(GRILLE, encoding="utf-8-sig")
import numpy as np
def minmax(x): return (x - x.min()) / (x.max() - x.min())
g["indice"] = g["score"].round(2)   # calculé par Scripts/score_snowballing.py (Google Scholar / Consensus calibré / imputé)
top = g[g["statut"] == "inclus"].sort_values("indice", ascending=False).head(N_TOP).reset_index(drop=True)

def initiales(prenom):
    p = [x for x in re.split(r"[ -]+", prenom) if x]
    return " ".join(x[0] + "." for x in p)
def formater_auteurs(a):
    auteurs = [re.sub(r"\s*\(dir\.\)", "", x.strip()) for x in str(a).split(";") if x.strip()]
    etal = "et al." in auteurs
    auteurs = [x for x in auteurs if x != "et al."]
    noms = []
    for x in auteurs:
        p = [y.strip() for y in x.split(",", 1)]
        noms.append(f"{p[0]}, {initiales(p[1])}" if len(p) == 2 and p[1] else p[0])
    if len(noms) == 1: out = noms[0]
    elif len(noms) == 2: out = " et ".join(noms)
    else: out = ", ".join(noms[:-1]) + " et " + noms[-1]
    return out + " et al." if etal else out
def ref_parts(r):
    lien = ""
    if isinstance(r["doi"], str) and r["doi"]: lien = " https://doi.org/" + r["doi"]
    elif isinstance(r["url"], str) and r["url"]: lien = " " + r["url"]
    return (f"{formater_auteurs(r['author'])} ({int(r['year'])}). {r['title']}. ", str(r["journal"]) if isinstance(r["journal"], str) else "", "." + lien)
refs = [ref_parts(r) for _, r in top.sort_values(["author", "year"]).iterrows()]

# ---------- 3. Document Word ----------
doc = Document()
for section in doc.sections:
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(2.5)

def set_font(run, size=12, bold=None, italic=None, upper=False):
    run.font.name = "Times New Roman"; run.font.size = Pt(size); run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr(); rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    for k in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"): rfonts.set(qn(k), "Times New Roman")
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if upper: run.text = run.text.upper()

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.paragraph_format.line_spacing = 2.0
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for name in ("Heading 1", "Heading 2", "Caption", "List Bullet"):
    st = doc.styles[name]
    st.font.name = "Times New Roman"; st.font.size = Pt(12); st.font.color.rgb = RGBColor(0, 0, 0)
    st.font.bold = name.startswith("Heading"); st.font.italic = False
    st.paragraph_format.line_spacing = 2.0; st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if name != "Caption" else WD_ALIGN_PARAGRAPH.LEFT
    if st.element.rPr is not None and st.element.rPr.rFonts is not None:
        st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman"); st.element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman"); st.element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")

CITES = set()
AUT = dict(zip(g["id"], g["citationAuteurAnnee"].str.strip("()")))
AUT["jalali2012"] = "Jalali et Wohlin, 2012"
def resoudre_citations(text):
    """[@a; @b] -> (A, année ; B, année) ; @a -> A (année)"""
    def bloc(m):
        cles = re.findall(r"@([A-Za-z0-9_]+)", m.group(0)); CITES.update(cles)
        return "(" + " ; ".join(AUT.get(k, k) for k in cles) + ")"
    text = re.sub(r"\[@[^\]]+\]", bloc, text)
    def libre(m):
        k = m.group(1); CITES.add(k); a = AUT.get(k, k)
        return re.sub(r", (\d{4})$", r" (\1)", a)
    return re.sub(r"(?<![\w\[])@([A-Za-z0-9_]+)", libre, text)

def add_inline(p, text):
    """gras **x**, italique *x* ; liens laissés en clair ; justification et interligne double explicites"""
    text = resoudre_citations(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    for tok in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
        if not tok: continue
        if tok.startswith("**"): r = p.add_run(tok[2:-2]); set_font(r, bold=True)
        elif tok.startswith("*"): r = p.add_run(tok[1:-1]); set_font(r, italic=True)
        elif tok.startswith("`"): r = p.add_run(tok[1:-1]); set_font(r)
        else:
            tok = re.sub(r"<(https?://[^>]+)>", r"\1", tok)
            r = p.add_run(tok); set_font(r)

def add_page_field(paragraph):
    run = paragraph.add_run(); set_font(run)
    for tag, text in (("begin", None), (None, "PAGE"), ("end", None)):
        if tag:
            el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), tag)
        else:
            el = OxmlElement("w:instrText"); el.set(qn("xml:space"), "preserve"); el.text = text
        run._r.append(el)

# En-tête (identifiants, gras, majuscules, à droite) sur toutes les pages
hdr = doc.sections[0].header.paragraphs[0]
hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = hdr.add_run(IDUL); set_font(r, bold=True, upper=True)
hdr.paragraph_format.line_spacing = 1.0

# ---- Page titre (section 1, sans numéro de page) ----
body_words = 0
def count_words(t): return len([w for w in re.split(r"\s+", t) if re.search(r"[0-9A-Za-zÀ-ÿ]", w)])
for b in blocks:
    if b[0] == "bibliographie": break
    if b[0] in ("p", "li") or b[0].startswith("h") or b[0].startswith("figure"): body_words += count_words(b[1])
    if b[0] == "tableau": body_words += 0

def centered(text, size=12, bold=False, space_before=0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.line_spacing = 2.0
    r = p.add_run(text); set_font(r, size=size, bold=bold); return p

centered("", space_before=72)
centered("TITRE", size=16, bold=True)
centered("", space_before=24)
centered("Question de recherche : entre 1961 et 2021, comment l'écart de revenu d'emploi et de statut socioprofessionnel entre francophones et anglophones de langue maternelle au Québec a-t-il évolué, et à partir de quel recensement cet écart s'est-il inversé, si tant est qu'il l'ait fait ?")
centered("", space_before=48)
centered("[Sigle et titre du cours]")
centered("[Département et université]")
centered("", space_before=24)
mois = ["janvier","février","mars","avril","mai","juin","juillet","août","septembre","octobre","novembre","décembre"]
d = datetime.date.today(); centered(f"{d.day} {mois[d.month-1]} {d.year}")
centered("", space_before=24)
centered(f"Nombre total de mots : {body_words} (page titre et bibliographie exclues)")

# ---- Section 2 : corps, numéroté à partir de 1 ----
sec = doc.add_section(WD_SECTION.NEW_PAGE)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.5)
sec.header.is_linked_to_previous = True
sec.footer.is_linked_to_previous = False
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_page_field(fp)
pg = OxmlElement("w:pgNumType"); pg.set(qn("w:start"), "1")
cols = sec._sectPr.find(qn("w:cols"))          # ordre du schéma : pgNumType avant cols
if cols is not None: cols.addprevious(pg)
else: sec._sectPr.append(pg)

h1 = h2 = nfig = 0
for kind, content in blocks:
    if kind == "h1" and content.startswith("Bibliographie"):
        p = doc.add_paragraph(style="Heading 1"); r = p.add_run("Bibliographie"); set_font(r, bold=True)
        cle_top = set(top["id"])
        sel = g[g["id"].isin(cle_top | CITES)].sort_values(["author", "year"])
        for _, r_ in sel.iterrows():
            a, j, c = ref_parts(r_)
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.25); p.paragraph_format.first_line_indent = Cm(-1.25)
            r = p.add_run(a); set_font(r)
            if j: r = p.add_run(j); set_font(r, italic=True)
            r = p.add_run(c); set_font(r)
        if "jalali2012" in CITES:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.25); p.paragraph_format.first_line_indent = Cm(-1.25)
            r = p.add_run("Jalali, S. et Wohlin, C. (2012). Systematic literature studies: Database searches vs. backward snowballing. "); set_font(r)
            r = p.add_run("Proceedings of the ACM-IEEE International Symposium on Empirical Software Engineering and Measurement"); set_font(r, italic=True)
            r = p.add_run(", 29-38. https://doi.org/10.1145/2372251.2372257"); set_font(r)
    elif kind == "h1":
        h1 += 1; h2 = 0
        p = doc.add_paragraph(style="Heading 1"); r = p.add_run(f"{h1}. {content}"); set_font(r, bold=True)
    elif kind == "h2":
        h2 += 1
        p = doc.add_paragraph(style="Heading 2"); r = p.add_run(f"{h1}.{h2} {content}"); set_font(r, bold=True)
    elif kind == "p":
        p = doc.add_paragraph(); add_inline(p, content)
    elif kind == "li":
        p = doc.add_paragraph(style="List Bullet"); add_inline(p, content)
    elif kind.startswith("figure"):
        nfig += 1
        img = PNG if kind.endswith("graphique") else PNG_RESEAU
        doc.add_page_break()                     # une figure par page
        doc.add_picture(img, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(style="Caption"); r = p.add_run(f"Figure {nfig}. " + content); set_font(r)
        p.paragraph_format.line_spacing = 1.0
    elif kind == "tableau":
        t = doc.add_table(rows=1, cols=6); t.style = "Table Grid"
        for i, h in enumerate(["Rang", "Publication", "Citations", "Source", "Récurrence", "Indice"]):
            c = t.rows[0].cells[i]; c.text = ""; r = c.paragraphs[0].add_run(h); set_font(r, size=10, bold=True)
        for i, row in top.iterrows():
            cells = t.add_row().cells
            src = {"Google Scholar": "GS", "imputé (médiane)": "imputé"}.get(str(row.get("citations_score_source", "")), "Cons. cal.")
            vals = [str(i + 1), row["citationAuteurAnnee"], str(int(row["citations_score"])), src, str(int(row["recurrence"])), f"{row['indice']:.2f}"]
            for c, v in zip(cells, vals):
                c.text = ""; r = c.paragraphs[0].add_run(v); set_font(r, size=10)
                c.paragraphs[0].paragraph_format.line_spacing = 1.0
        p = doc.add_paragraph(style="Caption"); r = p.add_run("Tableau 1. Composantes de l'indice de pertinence des trente lectures (GS : Google Scholar ; Cons. cal. : compte Consensus ramené à l'échelle Google Scholar ; imputé : médiane)."); set_font(r)
        p.paragraph_format.line_spacing = 1.0
    elif kind == "bibliographie" or (kind == "h1" and content == "Bibliographie" and False):
        for a, j, c in refs:
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.25); p.paragraph_format.first_line_indent = Cm(-1.25)
            r = p.add_run(a); set_font(r)
            if j: r = p.add_run(j); set_font(r, italic=True)
            r = p.add_run(c); set_font(r)

doc.save(OUT)
# correctif de schéma : python-docx omet l'attribut w:percent de <w:zoom> dans settings.xml
import zipfile, shutil
tmp = OUT + ".tmp"
with zipfile.ZipFile(OUT) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "word/settings.xml":
            data = re.sub(r"<w:zoom(?![^>]*w:percent)([^>]*)/>", r'<w:zoom w:percent="100"\1/>', data.decode("utf-8")).encode("utf-8")
        zout.writestr(item, data)
shutil.move(tmp, OUT)
print("docx écrit :", OUT, "| mots du corps :", body_words, "| blocs :", len(blocks))
